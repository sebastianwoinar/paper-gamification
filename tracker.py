import os
import sys
import time
import operator
import logging
import requests
import json
import re
import docx

import detex

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from watchdog.events import FileModifiedEvent

class GamificationHandler(FileSystemEventHandler):


	def __init__(self, paper_filename, publish_url, paper_id):

		FileSystemEventHandler.__init__(self)

		self.paper_filename = paper_filename
		self.publish_url = publish_url
		self.paper_id = paper_id

		logging.info("Creating a GamificationHandler with paper: " +
			paper_filename +
			" publish_url: " +
			publish_url +
			" and paper id: " +
			paper_id
		)

		self.stats = {}
		self.words = {}
		self.paragraphs = []
		self.num_words = 0
		self.total_word_len = 0


	def on_modified(self, event):
		# MAIN CALLBACK - a file got modified
		logging.info("Modify event occurred: " + event.src_path)
		paper_path = os.path.abspath(self.paper_filename)

		if paper_path in event.src_path:                    
			logging.info("Paper change detected, calculating statistics ...")
			if type(event) == FileModifiedEvent :
				logging.info("A file was modified: " + event.src_path)
			else:
				logging.info("Following Event Type" + str(type(event)))

			self.calculate_statistics()
			logging.info("Publishing ...")
			self.publish()
			logging.info("Published!")


	def parse_paragraphs(self, text):
		# Will only work for markdown elements
		# 	divided by '##' markers
		oldline = ""
		for line in text.split('\n'):
			if line.startswith('## '):
				if oldline != "":
					# Count previous paragraph
					paragraph = text.split(oldline)[1].split(line)[0]
					self.count_paragraph_words(oldline, paragraph)
				oldline = line
		# Count last paragraph
		if oldline != "":
			paragraph = text.split(oldline)[1]
			self.count_paragraph_words(oldline, paragraph)


	def count_paragraph_words(self, line, paragraph):
		num_words = len(re.findall(r"[\w']+", paragraph))
		self.paragraphs.append((line.replace('#', '').strip(), num_words))


	def parse_text_statistics(self, text):
		for w in text:
			word = w.strip().lower()

			# Add to total_word_len 
			# to determine average word length later
			self.total_word_len += len(word)

			# Count distinct words with occurrences
			if word not in self.words:
				self.words[word] = 0
			self.words[word] += 1

			# Count all words
			self.num_words += 1


	def parse_word_file(self, filename):
		# Read file
		document = docx.opendocx(filename)
		text = " ".join(docx.getdocumenttext(document))
		self.parse_paragraphs(text)
		word_split = re.findall(r"[\w']+", text)

		# Analyse
		self.parse_text_statistics(word_split)


	def parse_text_file(self, filename):
		# Read file
		f = open(filename)

		text = ""
		for line in f.readlines():
			text += line
			word_split = re.findall(r"[\w']+", line)
			# Analyse
			self.parse_text_statistics(word_split)

		f.close()
		self.parse_paragraphs(text)

	def parse_tex_file(self, filename):
		# Read file
		f = open(filename)

		text = ""
		for line in f.readlines():
			clean_line = detex.detex(line)
			
			text += clean_line
			word_split = re.findall(r"[\w']+", clean_line)
			# Analyse
			self.parse_text_statistics(word_split)

		f.close()
		self.parse_paragraphs(text)
		self.exclude_words("./tex_commands.txt")

	def exclude_words(self, filename):
		words = []
		num_words = 0
		# Count	and compare
		f = open(filename)
		for word in f.readlines():
			if word.strip() != "":
				word_to_exclude = word.strip().lower()
				if self.words.has_key(word_to_exclude):
					del self.words[word_to_exclude]
		f.close()

	def parse_file(self, filename):
		# Parse file 

		if filename.endswith(".docx"):
			logging.info("\tParsing the Word document " + filename)
			self.parse_word_file(filename)
		elif  filename.endswith(".tex"):
			logging.info("\tParsing the tex file " + filename)
			self.parse_tex_file(filename)
		elif filename.endswith(".txt") or filename.endswith(".tex") or filename.endswith(".md"):
			logging.info("\tParsing the file " + filename)
			self.parse_text_file(filename)
            
	def calculate_statistics(self):
		# Reset values 
		self.stats = {}
		self.words = {}
		self.paragraphs = []
		self.num_words = 0
		self.total_word_len = 0

		if os.path.isdir(self.paper_filename):
			logging.info("\tParsing all files in directory ...")
			path = self.paper_filename
			if not path.endswith("/"):
				path = path + "/"
				
			dirList = os.listdir(path)
			for entry in dirList:
				if os.path.isdir(path + entry) == True:
					logging.info("\tFound subdirectory - but there has to be an end ;-)")                    
				else:
					self.parse_file( path + entry )
		else:
			self.parse_file(self.paper_filename)

		# By now, text-statistics should be saved in instance variables
		
		self.build_stats()

		logging.info("\tStats: " + str(self.stats))

	def build_stats(self):
		# Build stats together
		# Determine interesting words
		logging.info("\tCalculating interesting words ...")
		interesting_words = self.get_interesting_words(40)

		# Determine average word length 
		logging.info("\tCalculating average word length ...")
		avg_len = float(self.total_word_len) / float(self.num_words)

		# Determine Oxford coverage
		logging.info("\tCalculating oxford coverage ...")
		oxford_coverage = self.get_coverage("./oxford.txt")

		# Determine Fancy word coverage
		logging.info("\tCalculating fancy words coverage ...")
		fancy_coverage = self.get_coverage("./fancy.txt")

		# Determine academic word list coverage 
		logging.info("\tCalculating academic word list coverage ...")
		awl_coverage = self.get_awl_coverage("./awl.txt")

		logging.info("\tBuilding stats together ...")
		self.stats = {
			"num_words" : self.num_words,
			"different_words" : len(self.words),
			"avg_len" : avg_len,
			"paragraphs": self.paragraphs,
			"interesting_words": interesting_words,
			"oxford_coverage" : {
				"total" : oxford_coverage["total"],
				"num_hits": len(oxford_coverage["hits"])
			},
			"fancy_coverage" : {
				"total" : fancy_coverage["total"],
				"num_hits": len(fancy_coverage["hits"])
			},
			"awl_coverage" : {
				"words_total": awl_coverage["words_total"],
				"words_hits": awl_coverage["words_hits"],
				"category_total": awl_coverage["category_total"],
				"category_num_hits": awl_coverage["category_num_hits"],
				"category_hits": awl_coverage["category_hits"]
			}
		}

	def get_interesting_words(self, num):
		sorted_words = sorted(self.words.iteritems(), key=operator.itemgetter(1), reverse=True)
		interesting_words = []

		num = min(num, len(sorted_words))
		min_len = 10

		while len(interesting_words) != num: # As long as we don't have as many words as we want
			for word in sorted_words:
				if len(word[0]) >= min_len:
					if word[1] == 1:
						# Word only occurs once in the text
						# -> since sorted_words is sorted by occurrence:
						#    break and go down with min word length
						break
					if word not in interesting_words:
						interesting_words.append(word)
				if len(interesting_words) == num:
					# Got enough words, break will break both loops
					break
			min_len -= 1
			if min_len < 2: # Text contains really few words, we just have to add them, until we have enough
				for word in sorted_words:
					if word not in interesting_words:
						interesting_words.append(word)
						if len(interesting_words) == num:
							# Got enough words, break will break both loops
							break

		# Sort result and return
		interesting_words = sorted(interesting_words, key=operator.itemgetter(1), reverse=True)
		return interesting_words


	
	def get_coverage(self, filename):
		""" Reads a list of words and compares it to the own words"""
		words = []
		num_words = 0
		# Count	and compare
		f = open(filename)
		for word in f.readlines():
			if word.strip() != "":
				words.append(word.strip().lower())
				num_words += 1

		hits = set(words).intersection(set(self.words.keys()))
		f.close()
		return { "total": num_words, "hits": list(hits)}


	def get_awl_coverage(self, filename):
		words = {}
		f = open(filename)

		category = ""
		for word in f.readlines():
			if not word.startswith('\t'):
				category = word.strip()
			words[word.strip()] = category

		hits = set(words.keys()).intersection(set(self.words.keys()))

		category_hits = {}
		for category in set(words.values()):
			category_hits[category] = 0

		for hit in hits:
			category_hits[words[hit]] += 1

		category_num_hits = 0
		for key in category_hits.keys():
			if category_hits[key] > 0:
				category_num_hits += 1

		return {
			"words_total": len(words),
			"words_hits": len(list(hits)),
			"category_total": len(list(set(words.values()))),
			"category_num_hits": category_num_hits,
			"category_hits": category_hits
		}


	def publish(self):
		payload = {"stats" : json.dumps(self.stats)}
		r = requests.put(self.publish_url + "/papers/" + self.paper_id + ".json", data=payload)


def set_paper_alive(publish_url, paper_id, alive):
	payload = {"alive" : str(alive).lower()}
	r = requests.put(publish_url + "/papers/" + paper_id + ".json", params=payload)


if __name__ == "__main__":
	logging.basicConfig(level=logging.INFO,
						format='%(asctime)s - %(message)s',
						datefmt='%Y-%m-%d %H:%M:%S')
	if len(sys.argv) != 4:
		print "Usage: python tracker.py <paper-file> <publish-host> <paper-id>"
		sys.exit()

	# Parse command line params
	filename = sys.argv[1]
	publish_url = sys.argv[2]
	paper_id = sys.argv[3]

	path = os.path.dirname(os.path.abspath(filename))

	# Enable "Currently writing..."
	set_paper_alive(publish_url, paper_id, True);

	# Observer setup
	event_handler = GamificationHandler(filename, publish_url, paper_id)
	observer = Observer()
	logging.info("Starting observer with watch path: " + path)
	observer.schedule(event_handler, path=path, recursive=True)
	# Observer start 
	observer.start()
	logging.info("Observer started.")

	try:
		while True:
			time.sleep(1)
	except KeyboardInterrupt:
		# Disable "Currently writing..."
		set_paper_alive(publish_url, paper_id, False);
		observer.stop()
	observer.join()
